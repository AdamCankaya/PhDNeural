"""Generate IEEE-style outline DOCX for BRCA TCGA NAS vs Bio-NAS paper."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def set_run_font(run, name="Times New Roman", size=10, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def style_paragraph(p, space_before=0, space_after=6, line_spacing=1.08, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if align is not None:
        p.alignment = align


def add_centered(doc, text, size=10, bold=False, italic=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    style_paragraph(
        p, space_before=space_before, space_after=space_after, align=WD_ALIGN_PARAGRAPH.CENTER
    )
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_body(doc, text, size=10, italic=False, space_before=0, space_after=6, indent=False):
    p = doc.add_paragraph()
    style_paragraph(p, space_before=space_before, space_after=space_after)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.2)
    run = p.add_run(text)
    set_run_font(run, size=size, italic=italic)
    return p


def add_heading_ieee(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        style_paragraph(p, space_before=12, space_after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(text.upper())
        set_run_font(run, size=10, bold=True)
    elif level == 2:
        style_paragraph(p, space_before=10, space_after=4)
        run = p.add_run(text)
        set_run_font(run, size=10, bold=True)
    elif level == 3:
        style_paragraph(p, space_before=8, space_after=3)
        run = p.add_run(text)
        set_run_font(run, size=10, bold=True, italic=True)
    else:
        style_paragraph(p, space_before=6, space_after=3)
        run = p.add_run(text)
        set_run_font(run, size=10, italic=True)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    style_paragraph(p, space_before=0, space_after=2)
    p.paragraph_format.left_indent = Inches(0.25 + 0.2 * level)
    run = p.add_run(text)
    set_run_font(run, size=10)
    return p


def add_note(doc, text):
    return add_body(doc, text, size=9, italic=True, space_after=8)


def build_document(out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)

    # Front matter
    add_centered(
        doc,
        "Comparative Spatio-Temporal Neural Architecture Search versus "
        "Biologically-Informed NAS for Multi-Omic Classification on TCGA-BRCA",
        size=14,
        bold=True,
        space_after=10,
    )
    add_centered(doc, "Adam Cankaya", size=11, space_after=2)
    add_centered(
        doc,
        "[Affiliation, Department, Institution — TBD]",
        size=10,
        italic=True,
        space_after=2,
    )
    add_centered(doc, "[City, Country]  ·  [email@institution.edu]", size=9, space_after=8)
    add_centered(
        doc,
        "IEEE Manuscript Outline — Structure Only (Draft Scaffold)",
        size=9,
        italic=True,
        space_after=4,
    )
    add_note(
        doc,
        "Document status: outline for peer-reviewed CS / computational biology journal submission. "
        "Placeholder summary lines (italic) indicate intended content; replace with full prose, "
        "equations, tables, and figures prior to submission. Formatting approximates IEEE journal "
        "manuscript prep (Times New Roman; numbered sections). Target venue TBD.",
    )

    add_heading_ieee(doc, "Index Terms", level=2)
    add_body(
        doc,
        "Neural architecture search, biologically-informed deep learning, multi-omics, "
        "intermediate fusion, spatio-temporal learning, multi-task classification, Optuna, "
        "TCGA-BRCA, breast invasive carcinoma, tumor staging, KEGG, Reactome, Docker, "
        "hyperparameter optimization.",
        space_after=10,
    )

    # Abstract
    add_heading_ieee(doc, "Abstract", level=1)
    add_note(
        doc,
        "[~150–250 words when written.] BRCA-only study comparing unconstrained NAS (Track A) "
        "vs Bio-NAS (Track B; KEGG and Reactome pathway constraints) for multi-task classification "
        "of TCGA clinical labels—patient disease status (phenotype from sample_type) and AJCC "
        "pathologic tumor stage (severity). Methods use spatio-temporal multi-omic tensors, "
        "Intermediate Fusion, phased Optuna, patient-level holdout, and Docker-reproducible "
        "pipelines.",
    )
    add_body(
        doc,
        "Placeholder abstract: This paper compares unconstrained Neural Architecture Search (NAS) "
        "and Biologically-Informed NAS (Bio-NAS) on The Cancer Genome Atlas Breast Invasive "
        "Carcinoma (TCGA-BRCA) cohort only. Using spatio-temporal multi-omic inputs "
        "(methylation and RNA-seq), we perform multi-task classification of disease status and "
        "tumor stage labels derived from TCGA clinical fields. … [complete after results.]",
        indent=True,
    )

    # I. Introduction
    add_heading_ieee(doc, "I. Introduction", level=1)
    add_note(
        doc,
        "Motivate BRCA multi-omic classification from TCGA labels; limitations of hand-designed "
        "nets and early raw-concat fusion; introduce spatio-temporal NAS vs Bio-NAS; state "
        "BRCA-only scope, contributions, and organization.",
    )

    add_heading_ieee(doc, "A. Motivation and Clinical / Computational Context", level=2)
    add_bullet(doc, "BRCA as the sole cohort for this manuscript (vertical-slice study).")
    add_bullet(
        doc,
        "Clinical classification targets aligned with TCGA fields: disease status and tumor stage.",
    )
    add_bullet(doc, "Curse of dimensionality and modality dominance under early fusion.")
    add_bullet(doc, "Need for reproducible, leakage-safe patient-level evaluation (Docker-first).")

    add_heading_ieee(doc, "B. Neural Architecture Search (NAS) in Biomedical ML", level=2)
    add_bullet(doc, "Brief survey of NAS for tabular / omic / spatio-temporal settings (citations TBD).")
    add_bullet(doc, "Role of hyperparameter and architecture search with Optuna / Hyperband.")

    add_heading_ieee(doc, "C. Biologically-Informed Learning and Pathway Priors", level=2)
    add_bullet(
        doc,
        "Gene regulatory / pathway graphs from both KEGG and Reactome as Track B constraints.",
    )
    add_bullet(
        doc,
        "Hypothesis: Track B (Bio-NAS) improves generalization and/or interpretability vs Track A "
        "under matched spatio-temporal Intermediate Fusion and compute budget.",
    )

    add_heading_ieee(doc, "D. Contributions", level=2)
    add_bullet(
        doc,
        "BRCA-only dual-track A/B protocol for multi-task TCGA label classification "
        "(disease status + AJCC pathologic tumor stage).",
    )
    add_bullet(
        doc,
        "Spatio-temporal Intermediate Fusion search space (MethEncoder, RNAEncoder, temporal "
        "modules, FusionDecoder).",
    )
    add_bullet(
        doc,
        "Phased Optuna search with Track B masks from KEGG and Reactome; distributed study storage.",
    )
    add_bullet(
        doc,
        "Docker-reproducible data and experiment workflow for open-access GDC/TCGA artifacts.",
    )

    add_heading_ieee(doc, "E. Paper Organization", level=2)
    add_note(doc, "One short paragraph mapping Sections II–X.")

    # II. Related Work
    add_heading_ieee(doc, "II. Related Work", level=1)
    add_note(
        doc,
        "Position against prior multi-omic fusion, staging/phenotype classifiers, NAS, and "
        "biology-informed DL on TCGA-BRCA.",
    )

    add_heading_ieee(doc, "A. Multi-Omic Integration and Fusion Strategies", level=2)
    add_bullet(doc, "Early, intermediate, and late fusion taxonomies.")
    add_bullet(doc, "Prior BRCA / TCGA multi-omic classification models (citations TBD).")

    add_heading_ieee(doc, "B. Spatio-Temporal Deep Learning for Genomics", level=2)
    add_bullet(doc, "Spatial genomic structure (CpG/gene neighborhoods, chromosomal spacing).")
    add_bullet(doc, "Temporal / longitudinal progression and irregular Δt modeling.")

    add_heading_ieee(doc, "C. Neural Architecture Search for Healthcare and Genomics", level=2)
    add_bullet(doc, "AutoML/NAS applied to clinical and molecular data.")
    add_bullet(doc, "Search algorithms: Bayesian optimization, Hyperband, evolutionary NAS.")

    add_heading_ieee(doc, "D. Biologically Constrained Neural Networks", level=2)
    add_bullet(doc, "Pathway-masked layers; priors from KEGG and Reactome.")
    add_bullet(
        doc,
        "Gap: controlled A/B of unconstrained vs pathway-constrained NAS under identical "
        "spatio-temporal fusion, labels, and BRCA cohort.",
    )

    # III. Background
    add_heading_ieee(doc, "III. Background: NAS versus Bio-NAS", level=1)
    add_note(
        doc,
        "Define Track A (control) vs Track B (innovation) without conflating Intermediate Fusion "
        "or spatio-temporal modules with biological constraints.",
    )

    add_heading_ieee(doc, "A. Unconstrained Neural Architecture Search (Track A)", level=2)
    add_bullet(
        doc,
        "Search over spatial/temporal encoder families, widths/depths, regularization, fusion heads.",
    )
    add_bullet(doc, "No pathway adjacency mask; connectivity free within module families.")
    add_note(
        doc,
        "Summary: Track A optimizes multi-task classification performance without biological "
        "graph constraints.",
    )

    add_heading_ieee(doc, "B. Biologically-Informed NAS (Track B / Bio-NAS)", level=2)
    add_bullet(
        doc,
        "Same spatio-temporal Intermediate Fusion scaffold as Track A.",
    )
    add_bullet(
        doc,
        "Pathway / adjacency masks constructed from both KEGG and Reactome (compare and/or union "
        "strategies — detail in Methods).",
    )
    add_bullet(
        doc,
        "Constraints applied to spatial / branch modules; frozen mask construction; dual-track Optuna.",
    )
    add_note(
        doc,
        "Summary: Bio-NAS restricts or biases architecture choices using prior molecular network "
        "knowledge from KEGG and Reactome.",
    )

    add_heading_ieee(doc, "C. Comparative Framing and Fairness Criteria", level=2)
    add_bullet(
        doc,
        "Shared BRCA cohort, preprocessing, Intermediate Fusion family, CV folds, and compute budget.",
    )
    add_bullet(doc, "Differ in biological constraint track (A vs B) and pathway-source variants.")
    add_bullet(
        doc,
        "Endpoints: multi-task classification metrics (accuracy/AUROC/AUPRC/macro-F1 per head); "
        "sparsity; attribution stability.",
    )

    # IV. Data
    add_heading_ieee(doc, "IV. Data Sources and Cohort Construction", level=1)
    add_note(
        doc,
        "This manuscript is BRCA-only. No non-BRCA cohorts are evaluated here.",
    )

    add_heading_ieee(doc, "A. TCGA-BRCA and GDC Access", level=2)
    add_bullet(doc, "Project: TCGA Breast Invasive Carcinoma (TCGA-BRCA) exclusively.")
    add_bullet(doc, "Access path: NCI GDC Portal / API; Level-3 open-access matrices preferred.")
    add_bullet(doc, "Controlled-access BAM/raw out of scope unless explicitly required later.")
    add_note(
        doc,
        "Summary: Describe download provenance, versioning, and open vs controlled data boundaries.",
    )

    add_heading_ieee(doc, "B. Multi-Omic Modalities and Spatio-Temporal Representation", level=2)
    add_bullet(
        doc,
        "DNA methylation (beta values; mean-impute NaNs; no Z-score/log for meth branch).",
    )
    add_bullet(doc, "RNA-seq transcriptomics (log2(TPM+1) then Z-score for RNA branch).")
    add_bullet(
        doc,
        "Spatial axis: genomic / CpG / gene structural spacing and neighborhood structure.",
    )
    add_bullet(
        doc,
        "Temporal axis: longitudinal visit structure, primary/recurrent matching, irregular Δt "
        "embedding where available.",
    )
    add_note(
        doc,
        "Summary: Construct 4D (or equivalent) spatio-temporal tensors for Intermediate Fusion "
        "loaders; BRCA feature-map and spacing configs as specified in the Bio-NAS plans.",
    )

    add_heading_ieee(doc, "C. Classification Targets from TCGA Clinical Fields", level=2)
    add_bullet(
        doc,
        "Primary tasks are supervised classification using TCGA-provided clinical labels "
        "(multi-task learning).",
    )
    add_bullet(
        doc,
        "Disease status / phenotype head: from TCGA sample_type "
        "(e.g., Solid Tissue Normal → 0; Primary/Metastatic/Recurrent Tumor → 1).",
    )
    add_bullet(
        doc,
        "Tumor stage / severity head: from ajcc_pathologic_tumor_stage "
        "(Stage I/IA/IB → 0; II/IIA/IIB → 1; III/IIIA/IIIB/IIIC → 2; IV → 3).",
    )
    add_bullet(
        doc,
        "Missing severity policy: mask (e.g., severity = −1 for normal samples without stage).",
    )
    add_bullet(
        doc,
        "Additional TCGA clinical classification labels may be added if consistently available; "
        "stage + disease status are the minimum reported task set.",
    )
    add_note(
        doc,
        "Summary: Align label definitions with the project disease registry for BRCA; report "
        "class balances and missingness.",
    )

    add_heading_ieee(doc, "D. Cohort Filtering and Longitudinal Matching", level=2)
    add_bullet(doc, "Case inclusion/exclusion rules for BRCA.")
    add_bullet(doc, "Primary vs recurrent (or multi-timepoint) matching for temporal modeling.")
    add_bullet(doc, "Irregular time-step handling and Δt features.")

    add_heading_ieee(doc, "E. Train / Validation / Holdout Splits", level=2)
    add_bullet(doc, "Patient-level 80/20 train–holdout; no patient leakage across folds.")
    add_bullet(doc, "Variance masks and scalers fit on training pool only.")
    add_bullet(doc, "5-fold CV within train pool for Optuna objectives.")
    add_note(doc, "Summary: Leakage-safe split protocol shared by Track A and Track B.")

    add_heading_ieee(doc, "F. Storage Layout and Reproducibility Artifacts", level=2)
    add_bullet(
        doc,
        "HDF5 / tensor serialization for Intermediate Fusion batches "
        "(methylation_tensor, rna_tensor, labels, temporal fields).",
    )
    add_bullet(doc, "Manifests, demographics, experiment logs, Optuna study identifiers.")

    # V. Methods
    add_heading_ieee(doc, "V. Methods", level=1)
    add_note(
        doc,
        "Core technical section: multi-task formulation, spatio-temporal modules, fusion, "
        "Optuna, Docker reproducibility, evaluation.",
    )

    add_heading_ieee(doc, "A. Problem Formulation (Multi-Task Classification)", level=2)
    add_bullet(
        doc,
        "Input: spatio-temporal multi-omic sample (x_meth, x_rna[, Δt / visit context]).",
    )
    add_bullet(
        doc,
        "Outputs: (i) disease-status / phenotype class; (ii) AJCC pathologic tumor-stage class; "
        "shared trunk with task-specific heads (MTL).",
    )
    add_bullet(
        doc,
        "Objective: maximize joint validation classification performance under Track A or Track B "
        "constraints (state loss weighting).",
    )

    add_heading_ieee(doc, "B. Network Structure and Module Families", level=2)
    add_heading_ieee(doc, "1) Methylation Encoder (MethEncoder)", level=3)
    add_bullet(doc, "Spatial CNN / transformer / MLP candidates over CpG structure.")
    add_bullet(doc, "Searchable HPs: depth, width, dropout, normalization, etc.")

    add_heading_ieee(doc, "2) RNA Encoder (RNAEncoder)", level=3)
    add_bullet(doc, "Parallel branch with modality-appropriate input scaling.")
    add_bullet(doc, "Searchable HPs analogous to methylation branch.")

    add_heading_ieee(doc, "3) Temporal Progression Modules", level=3)
    add_bullet(doc, "Modules consuming visit order / Δt for longitudinal BRCA samples.")
    add_bullet(doc, "Searchable temporal HPs (recurrent, attention, or Δt-embedding variants).")

    add_heading_ieee(doc, "4) Fusion Decoder and Classification Heads (FusionDecoder)", level=3)
    add_bullet(doc, "Post-concatenation dense tower → phenotype head + stage head.")
    add_bullet(doc, "MTL losses with masking for missing stage labels.")

    add_heading_ieee(doc, "5) Track B Pathway / Adjacency Constraints (KEGG and Reactome)", level=3)
    add_bullet(doc, "Build and freeze adjacency / pathway masks from KEGG.")
    add_bullet(doc, "Build and freeze adjacency / pathway masks from Reactome.")
    add_bullet(
        doc,
        "Report how each source (and any combined mask) alters searchable connectivity vs Track A.",
    )

    add_heading_ieee(doc, "C. Intermediate Fusion of Multi-Omic Data", level=2)
    add_bullet(doc, "Branch-specific encoding → latent concatenation → post-fusion dense → MTL heads.")
    add_bullet(
        doc,
        "Rationale vs early raw-concat (legacy) and late stacked fusion (optional baseline; TBD).",
    )
    add_bullet(doc, "Diagram placeholder: Fig. 1 — Spatio-temporal Intermediate Fusion NAS scaffold.")
    add_note(
        doc,
        "Summary: Intermediate Fusion is the default NAS architecture; early fusion retained only "
        "as a possible legacy software baseline (inclusion TBD).",
    )

    add_heading_ieee(doc, "D. Optuna-Based Architecture and Hyperparameter Search", level=2)
    add_heading_ieee(doc, "1) Search Space Definition", level=3)
    add_bullet(doc, "Categorical and continuous parameters for spatial, temporal, and decoder modules.")
    add_bullet(doc, "Track A vs Track B differences (KEGG/Reactome mask variants).")

    add_heading_ieee(doc, "2) Phased Search Protocol", level=3)
    add_bullet(doc, "Phase A: optimize branch / temporal HPs (encoders).")
    add_bullet(doc, "Phase B: freeze selected branch configs; optimize post-fusion dense + heads.")
    add_note(
        doc,
        "Summary: Phasing reduces joint search dimensionality and modality interference.",
    )

    add_heading_ieee(doc, "3) Pruners, Samplers, and Distributed Storage", level=3)
    add_bullet(doc, "TPE / Hyperband (or configured sampler/pruner); PostgreSQL Optuna RDB storage.")
    add_bullet(doc, "Parallel workers (Slurm / Docker) and study naming conventions.")

    add_heading_ieee(doc, "4) Trial Objective and Early Stopping", level=3)
    add_bullet(
        doc,
        "CV multi-task classification metric (e.g., mean AUROC / macro-F1 across heads, with "
        "masked stage samples handled correctly).",
    )
    add_bullet(doc, "Compute budget: trials × epochs × folds (state planned budget).")

    add_heading_ieee(doc, "E. Training Procedure", level=2)
    add_bullet(doc, "Optimizer, learning-rate schedule, batch size, regularization.")
    add_bullet(doc, "MTL classification losses with severity masking.")
    add_bullet(doc, "Hardware and software stack (PyTorch, CUDA, pinned container image).")

    add_heading_ieee(doc, "F. Docker-First Reproducible Workflow", level=2)
    add_note(
        doc,
        "Executable experiment paths are designed to run inside the Bio-NAS Docker container "
        "(docker/Dockerfile, docker-compose.yml), not as host-only workflows.",
    )
    add_bullet(
        doc,
        "Compose contract: from Bio-NAS/, docker compose up --build builds the image, optionally "
        "downloads open-access TCGA-BRCA sample data via GDC API on first run, and invokes "
        "container entrypoint scripts.",
    )
    add_bullet(
        doc,
        "Host bind mount: ./data/tcga → /data/tcga (cohort outputs under /data/tcga/BRCA).",
    )
    add_bullet(
        doc,
        "Dependencies pinned in docker/requirements.txt; scripts extended via "
        "scripts/docker_entrypoint.py and related container-invoked tools.",
    )
    add_bullet(
        doc,
        "Full-cohort ETL / Optuna workers remain Docker-oriented; interactive DUA / controlled-access "
        "credential steps (if any) are documented as out-of-container exceptions.",
    )
    add_bullet(
        doc,
        "Reference: Bio-NAS/docker/README.md for the current compose contract and expected artifacts "
        "(manifest.json, demographics, experiment outputs).",
    )

    add_heading_ieee(doc, "G. Baselines and Ablations (TBD)", level=2)
    add_bullet(
        doc,
        "Baseline set not finalized: candidates include legacy early fusion, classical ML / late "
        "stacking, spatial-only vs spatio-temporal ablations.",
    )
    add_bullet(doc, "Required ablation: Track A vs Track B; KEGG vs Reactome (vs combined) masks.")
    add_bullet(doc, "Finalize baseline table before experiments freeze.")

    add_heading_ieee(doc, "H. Evaluation Protocol", level=2)
    add_bullet(doc, "Locked patient holdout; per-head and joint classification metrics + CIs.")
    add_bullet(doc, "Confusion matrices / calibration for phenotype and stage heads.")
    add_bullet(doc, "Interpretability: attribution / pathway enrichment for Track B (KEGG & Reactome).")
    add_bullet(doc, "Statistical comparison of Track A vs Track B (paired tests / bootstrap).")

    # VI. Experimental Setup
    add_heading_ieee(doc, "VI. Experimental Setup", level=1)
    add_note(
        doc,
        "Concrete settings reviewers need to reproduce: tables of HPs, budgets, seeds, Docker image tags.",
    )

    add_heading_ieee(doc, "A. Implementation and Reproducibility", level=2)
    add_bullet(doc, "Code repository (PhDNeural / Bio-NAS), container image tag, dependency pins.")
    add_bullet(doc, "Random seeds; deterministic flags where feasible.")
    add_bullet(doc, "Document exact docker compose / entrypoint commands used for reported runs.")

    add_heading_ieee(doc, "B. Compute Infrastructure", level=2)
    add_bullet(doc, "Worker nodes, GPUs, Optuna hub (PostgreSQL), job orchestration.")

    add_heading_ieee(doc, "C. Hyperparameter Ranges and Search Budgets", level=2)
    add_bullet(doc, "Table II (placeholder): search space ranges for Track A and Track B.")
    add_bullet(doc, "Table III (placeholder): trials, wall-clock, GPU-hours.")

    add_heading_ieee(doc, "D. Metrics and Reporting Conventions", level=2)
    add_bullet(
        doc,
        "Primary: holdout multi-task classification performance (state aggregation across heads).",
    )
    add_bullet(doc, "Secondary: per-head metrics, sparsity, pathway attribution overlap, runtime.")

    # VII. Results
    add_heading_ieee(doc, "VII. Results", level=1)
    add_note(doc, "[To be filled after experiments.] Structure reserved for peer-review completeness.")

    add_heading_ieee(doc, "A. Cohort and Label Statistics", level=2)
    add_bullet(
        doc,
        "Table/figure: BRCA sample counts, modality coverage, phenotype and stage class balances, "
        "missing stage rates.",
    )

    add_heading_ieee(doc, "B. Search Dynamics and Selected Architectures", level=2)
    add_bullet(doc, "Optuna curves; best Trial A vs B architectures (and KEGG/Reactome variants).")
    add_bullet(doc, "Fig. 2 (placeholder): trial objective vs trial number.")

    add_heading_ieee(doc, "C. Classification Performance: Track A (NAS) vs Track B (Bio-NAS)", level=2)
    add_bullet(doc, "CV and holdout tables for phenotype and stage heads; confidence intervals.")
    add_bullet(doc, "Head-to-head comparison under matched budgets.")

    add_heading_ieee(doc, "D. Ablation Studies", level=2)
    add_bullet(doc, "Spatial-only vs spatio-temporal; pathway source (KEGG / Reactome / both).")
    add_bullet(doc, "Fusion / baseline ablations as finalized in Section V-G.")

    add_heading_ieee(doc, "E. Efficiency, Sparsity, and Resource Trade-offs", level=2)
    add_bullet(doc, "Parameters, FLOPs, training time, effective sparsity under Bio-NAS masks.")

    add_heading_ieee(doc, "F. Interpretability and Biological Plausibility", level=2)
    add_bullet(
        doc,
        "Pathway-level attributions vs known BRCA biology; compare KEGG- vs Reactome-guided models.",
    )

    # VIII. Discussion
    add_heading_ieee(doc, "VIII. Discussion", level=1)

    add_heading_ieee(doc, "A. Interpretation of NAS vs Bio-NAS Findings", level=2)
    add_note(
        doc,
        "Discuss whether biological priors helped phenotype classification, stage classification, "
        "both, or traded accuracy for interpretability.",
    )

    add_heading_ieee(doc, "B. Spatio-Temporal and Intermediate Fusion Design Choices", level=2)
    add_note(
        doc,
        "Reflect on temporal signal strength in BRCA, modality imbalance, and phased Optuna effectiveness.",
    )

    add_heading_ieee(doc, "C. Limitations", level=2)
    add_bullet(doc, "BRCA-only scope; no cross-disease claims in this paper.")
    add_bullet(doc, "TCGA clinical label noise / staging incompleteness; open-access Level-3 limits.")
    add_bullet(doc, "Incomplete pathway coverage; KEGG vs Reactome mapping disagreements.")
    add_bullet(doc, "Search budget and stochasticity of NAS; baseline set may evolve.")

    add_heading_ieee(doc, "D. Threats to Validity", level=2)
    add_bullet(doc, "Internal: leakage controls, multiple-testing, hyperparameter overfitting to CV.")
    add_bullet(doc, "External: other cancers / non-TCGA cohorts (future work).")
    add_bullet(doc, "Construct: whether stage/phenotype metrics capture clinical utility.")

    add_heading_ieee(doc, "E. Implications and Future Multi-Disease Extension", level=2)
    add_note(
        doc,
        "Bridge to broader Bio-NAS thesis agenda without overclaiming beyond BRCA results.",
    )

    # IX. Conclusion
    add_heading_ieee(doc, "IX. Conclusion", level=1)
    add_note(
        doc,
        "Restate BRCA-only multi-task classification problem, dual-track spatio-temporal method, "
        "main empirical takeaway (TBD), and contribution. Avoid new results.",
    )
    add_body(
        doc,
        "Placeholder conclusion: We will report a controlled comparison of unconstrained NAS and "
        "Bio-NAS (KEGG and Reactome) for spatio-temporal Intermediate Fusion multi-task "
        "classification of TCGA-BRCA disease status and tumor stage. …",
        indent=True,
    )

    # X. Summary
    add_heading_ieee(doc, "X. Summary", level=1)
    add_note(
        doc,
        "Short executive wrap-up. If the eventual venue disallows a separate Summary, merge into "
        "Conclusion or Abstract (venue TBD).",
    )
    add_bullet(doc, "Cohort: TCGA-BRCA only (GDC open-access multi-omics).")
    add_bullet(
        doc,
        "Tasks: classify disease status (phenotype) and AJCC pathologic tumor stage (severity).",
    )
    add_bullet(doc, "Inputs: spatial + temporal multi-omic tensors; Intermediate Fusion.")
    add_bullet(doc, "Search: phased Optuna; Track A (NAS) vs Track B (Bio-NAS; KEGG + Reactome).")
    add_bullet(doc, "Reproducibility: Docker Compose / pinned container workflow.")
    add_bullet(doc, "Take-home: [fill after results].")

    add_heading_ieee(doc, "Acknowledgment", level=1)
    add_note(doc, "Funding, compute sponsors, data acknowledgments (TCGA/GDC), and colleagues.")

    add_heading_ieee(doc, "References", level=1)
    add_note(
        doc,
        "IEEE numbered citation list. Seed topics: NAS surveys; Optuna; multi-omic fusion; "
        "spatio-temporal genomics; TCGA-BRCA staging; KEGG; Reactome; Hyperband; Docker "
        "reproducibility; pathway-informed DL.",
    )
    add_body(doc, "[1]  …", space_after=2)
    add_body(doc, "[2]  …", space_after=2)
    add_body(doc, "[3]  …", space_after=8)

    add_heading_ieee(doc, "Appendix A — Extended Search Space Tables", level=1)
    add_note(doc, "Full Optuna parameter catalogs for Track A and Track B (KEGG / Reactome).")

    add_heading_ieee(doc, "Appendix B — Additional Cohort, Label, and Preprocessing Details", level=1)
    add_note(
        doc,
        "GDC query filters; phenotype/stage mapping tables; QC plots; missingness handling.",
    )

    add_heading_ieee(doc, "Appendix C — Docker Reproduction Commands", level=1)
    add_note(
        doc,
        "Exact compose/build/run commands, volume mounts, image tags, and expected host-path outputs "
        "for reported experiments.",
    )

    add_heading_ieee(doc, "Appendix D — Supplementary Results", level=1)
    add_note(doc, "Per-fold metrics, failed trials, sensitivity analyses.")

    add_heading_ieee(doc, "Figure and Table Checklist (for manuscript completion)", level=1)
    add_bullet(doc, "Fig. 1: Spatio-temporal Intermediate Fusion (Track A vs Track B overlay).")
    add_bullet(doc, "Fig. 2: Optuna search dynamics.")
    add_bullet(doc, "Fig. 3: Holdout classification performance (phenotype + stage; NAS vs Bio-NAS).")
    add_bullet(doc, "Fig. 4: Ablation / sparsity / KEGG–Reactome attribution panel.")
    add_bullet(doc, "Table I: BRCA cohort, modality, and label statistics.")
    add_bullet(doc, "Table II: Search space summary.")
    add_bullet(doc, "Table III: Compute budget, Docker image tag, selected hyperparameters.")
    add_bullet(doc, "Table IV: Main classification results (CV + holdout).")
    add_bullet(doc, "Table V: Ablations and baselines (finalize when baseline set is chosen).")

    add_heading_ieee(doc, "Author Notes — Locked Manuscript Decisions", level=1)
    add_note(doc, "Decisions captured for drafting; remaining open items listed last.")
    add_bullet(doc, "Author: Adam Cankaya.")
    add_bullet(
        doc,
        "Primary tasks: multi-task classification of TCGA clinical labels — disease status "
        "(phenotype / sample_type) and AJCC pathologic tumor stage (severity).",
    )
    add_bullet(doc, "Cohort scope: BRCA only.")
    add_bullet(doc, "Data dimensionality: spatial and temporal.")
    add_bullet(doc, "Track B pathway sources: both KEGG and Reactome.")
    add_bullet(doc, "Reproducibility: Docker-first workflow documented in Methods + Appendix C.")
    add_bullet(doc, "Open: baselines set (early fusion / classical / late stacking) — TBD.")
    add_bullet(doc, "Open: target venue / whether separate Summary section is retained — TBD.")
    add_bullet(doc, "Open: affiliation block and corresponding email — TBD.")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    out_path = Path(__file__).resolve().parents[1] / "papers" / (
        "BRCA_TCGA_NAS_vs_Bio-NAS_IEEE_Outline.docx"
    )
    build_document(out_path)
    print(out_path)


if __name__ == "__main__":
    main()
